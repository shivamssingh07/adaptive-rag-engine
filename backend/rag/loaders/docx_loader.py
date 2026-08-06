"""DOCX document loader, backed by `python-docx`.

Extracts both paragraph text and table content (flattened row-by-row),
since Word documents frequently carry meaningful data in tables that a
naive paragraph-only extraction would silently drop.
"""

from __future__ import annotations

import logging
from pathlib import Path

import docx
from langchain_core.documents import Document

from backend.core.exceptions import CorruptedFileError
from backend.utils.ids import generate_document_id

logger = logging.getLogger(__name__)


class DOCXLoader:
    """Loads a `.docx` file into a single `Document`."""

    def load(self, file_path: Path) -> list[Document]:
        """Extract paragraph and table text from a Word document.

        Args:
            file_path: Path to the `.docx` file.

        Returns:
            A single-element list containing one `Document` with the full
            extracted text.

        Raises:
            CorruptedFileError: If the file cannot be opened as a valid
                `.docx` archive, or contains no extractable text.
        """
        document_id = generate_document_id()

        try:
            word_document = docx.Document(str(file_path))
        except Exception as exc:  # noqa: BLE001 - wrapped as a domain error
            raise CorruptedFileError(
                f"Could not open DOCX '{file_path.name}': {exc}",
                details={"file": file_path.name},
            ) from exc

        text_blocks: list[str] = [
            paragraph.text.strip()
            for paragraph in word_document.paragraphs
            if paragraph.text.strip()
        ]

        for table in word_document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_blocks.append(" | ".join(cells))

        full_text = "\n".join(text_blocks).strip()
        if not full_text:
            raise CorruptedFileError(
                f"DOCX '{file_path.name}' contains no extractable text.",
                details={"file": file_path.name},
            )

        logger.info("Parsed DOCX '%s': %d text block(s).", file_path.name, len(text_blocks))
        return [
            Document(
                page_content=full_text,
                metadata={
                    "source": file_path.name,
                    "document_id": document_id,
                    "file_type": "docx",
                },
            )
        ]
