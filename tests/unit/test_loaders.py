"""Unit tests for `backend.rag.loaders` — one format per test class."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest

from backend.core.exceptions import CorruptedFileError, UnsupportedFileTypeError
from backend.rag.loaders.csv_loader import CSVLoader
from backend.rag.loaders.docx_loader import DOCXLoader
from backend.rag.loaders.pdf_loader import PDFLoader
from backend.rag.loaders.registry import get_loader_for, supported_extensions
from backend.rag.loaders.text_loader import TextLoader


class TestPDFLoader:
    def test_loads_text_per_page(self, tmp_path: Path) -> None:
        pdf_path = tmp_path / "sample.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "This is page one about budgets.")
        doc.save(pdf_path)
        doc.close()

        documents = PDFLoader().load(pdf_path)

        assert len(documents) == 1
        assert "budgets" in documents[0].page_content.lower()
        assert documents[0].metadata["file_type"] == "pdf"
        assert documents[0].metadata["page"] == 1
        assert documents[0].metadata["source"] == "sample.pdf"

    def test_corrupted_pdf_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        bad_pdf = tmp_path / "bad.pdf"
        bad_pdf.write_bytes(b"not a real pdf")

        with pytest.raises(CorruptedFileError):
            PDFLoader().load(bad_pdf)


class TestDOCXLoader:
    def test_loads_paragraphs_and_tables(self, tmp_path: Path) -> None:
        import docx

        docx_path = tmp_path / "sample.docx"
        document = docx.Document()
        document.add_paragraph("This paragraph discusses quarterly revenue.")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Metric"
        table.cell(0, 1).text = "Value"
        document.save(docx_path)

        documents = DOCXLoader().load(docx_path)

        assert len(documents) == 1
        assert "revenue" in documents[0].page_content.lower()
        assert "Metric" in documents[0].page_content
        assert documents[0].metadata["file_type"] == "docx"

    def test_corrupted_docx_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        bad_docx = tmp_path / "bad.docx"
        bad_docx.write_bytes(b"not a real docx")

        with pytest.raises(CorruptedFileError):
            DOCXLoader().load(bad_docx)


class TestTextLoader:
    def test_loads_txt(self, tmp_path: Path) -> None:
        txt_path = tmp_path / "notes.txt"
        txt_path.write_text("This file discusses renewable energy.")

        documents = TextLoader().load(txt_path)

        assert len(documents) == 1
        assert documents[0].metadata["file_type"] == "text"
        assert "renewable" in documents[0].page_content.lower()

    def test_loads_markdown(self, tmp_path: Path) -> None:
        md_path = tmp_path / "readme.md"
        md_path.write_text("# Title\n\nDiscussion of wind turbines.")

        documents = TextLoader().load(md_path)

        assert documents[0].metadata["file_type"] == "markdown"
        assert "wind turbines" in documents[0].page_content.lower()

    def test_empty_file_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        empty_path = tmp_path / "empty.txt"
        empty_path.write_text("   \n  ")

        with pytest.raises(CorruptedFileError):
            TextLoader().load(empty_path)


class TestCSVLoader:
    def test_loads_one_document_per_row(self, tmp_path: Path) -> None:
        csv_path = tmp_path / "data.csv"
        csv_path.write_text("name,role\nAlice,Engineer\nBob,Manager\n")

        documents = CSVLoader().load(csv_path)

        assert len(documents) == 2
        assert "Alice" in documents[0].page_content
        assert documents[0].metadata["row"] == 1
        assert documents[1].metadata["row"] == 2

    def test_empty_csv_raises_corrupted_file_error(self, tmp_path: Path) -> None:
        csv_path = tmp_path / "empty.csv"
        csv_path.write_text("name,role\n")

        with pytest.raises(CorruptedFileError):
            CSVLoader().load(csv_path)


class TestLoaderRegistry:
    def test_supported_extensions(self) -> None:
        assert supported_extensions() == [".csv", ".docx", ".md", ".pdf", ".txt"]

    def test_dispatches_by_extension(self, tmp_path: Path) -> None:
        txt_path = tmp_path / "a.txt"
        txt_path.write_text("hello")
        assert isinstance(get_loader_for(txt_path), TextLoader)

    def test_unsupported_extension_raises(self, tmp_path: Path) -> None:
        bad_path = tmp_path / "a.png"
        bad_path.write_text("hello")
        with pytest.raises(UnsupportedFileTypeError):
            get_loader_for(bad_path)
